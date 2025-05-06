import subprocess
import socket
import ssl
import requests
import streamlit as st
import os
import io
import re
from docx import Document

# streamlit session init
if "docx_file" not in st.session_state:
    st.session_state.docx_file = None

# vulnerability mapping
nist_mappings = {
    "Apache/2.4.7": "CVE-2021-41773: Path traversal vulnerability in Apache HTTP Server 2.4.49.",
    "X-Frame-Options header is not present": "Potential clickjacking vulnerability.",
    "X-Content-Type-Options header is not set": "Risk of MIME type confusion attacks.",
    "mod_negotiation is enabled": "May allow attackers to brute force filenames (CVE-2009-1195)."
}

# functions
def run_command(command):
    try:
        result = subprocess.run(command, shell=True, capture_output=True, text=True, check=True)
        return result.stdout.strip()
    except subprocess.CalledProcessError as e:
        return e.output.strip() or str(e)

def run_nmap(target):
    return run_command(f"nmap -sS -sV --script vuln {target}")

def run_nikto(target):
    return run_command(f"nikto -h {target} -Tuning 1234567890abcde -maxtime 600 -timeout 20")

def check_ssl(target):
    context = ssl.create_default_context()
    try:
        with socket.create_connection((target, 443), timeout=5) as sock:
            with context.wrap_socket(sock, server_hostname=target) as ssock:
                cert = ssock.getpeercert()
                return (f"Issuer: {cert.get('issuer')}\nValid from: {cert.get('notBefore')}\nValid until: {cert.get('notAfter')}")
    except socket.gaierror:
        return "SSL check failed: Invalid hostname or DNS error."
    except ConnectionRefusedError:
        return "SSL check failed: Port 443 is closed on the target."
    except socket.timeout:
        return "SSL check failed: Connection timed out."
    except OSError as e:
        return f"SSL check failed: {e}"


def check_http_headers(target):
    try:
        response = requests.get(f"http://{target}", timeout=10, verify=True)
        headers = response.headers
        report = ""
        for header in ["Strict-Transport-Security", "Content-Security-Policy", "X-Frame-Options", "X-Content-Type-Options"]:
            report += f"{header}: {'Present' if header in headers else 'NOT present'}\n"
        return report
    except Exception as e:
        return f"Header check failed: {e}"

def simple_vulnerability_analysis(nmap_output, nikto_output, header_output):
    findings = []

    apache_match = re.search(r"Apache/2\.4\.\d+", nmap_output)
    if apache_match:
        version = apache_match.group(0)
        findings.append(f"{version} detected. Check for known vulnerabilities.")

    for line in header_output.splitlines():
        if "NOT present" in line:
            findings.append(f"Missing header: {line.split(':')[0]}. Potential risk.")

    if "mod_negotiation" in nikto_output:
        findings.append("mod_negotiation enabled. Might be vulnerable to CVE-2009-1195.")

    if "Server leaks inodes" in nikto_output:
        findings.append("Nikto detected inode information disclosure.")

    if "OSVDB-" in nikto_output:
        findings.append("Nikto reported vulnerabilities (OSVDB references found).")

    vuln_matches = re.findall(r"(CVE-\d{4}-\d{4,7})", nikto_output + nmap_output)
    for cve in set(vuln_matches):
        findings.append(f"Detected {cve}. Check NIST database for details.")

    if not findings:
        return "No critical vulnerabilities detected."

    return "\n".join([f"- {issue}" for issue in findings])


def save_report(nmap_output, nikto_output, ssl_output, header_output, analysis):
    try:
        doc = Document()
        doc.add_heading('Web Server Security Audit Report', 0)

        sections = [
            ("Nmap Scan Results", nmap_output),
            ("Nikto Scan Results", nikto_output),
            ("SSL Certificate Information", ssl_output),
            ("HTTP Security Header Check", header_output),
            ("Analysis and Vulnerability Summary", analysis),
            ("Recommended Actions", 
             "- Update outdated services.\n"
             "- Enable HTTP security headers.\n"
             "- Disable unnecessary modules.\n"
             "- Monitor CVEs and patch frequently.\n"
             "- Use strong SSL configurations.\n"
             "- Perform periodic security scans.")
        ]

        for title, content in sections:
            doc.add_heading(title, level=1)
            doc.add_paragraph(content)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    except Exception as e:
        st.error(f"Report generation failed: {e}")
        return None

# streamlit ui
st.title("Web Server Security Auditor")

target = st.text_input("Enter target domain or IP (without https://):")

if st.button("Run Audit") and target:
    with st.spinner("Running Nmap..."):
        st.session_state.nmap_result = run_nmap(target)
    st.text_area("Nmap Output", st.session_state.nmap_result, height=200)

    with st.spinner("Running Nikto..."):
        st.session_state.nikto_result = run_nikto(target)
    st.text_area("Nikto Output", st.session_state.nikto_result, height=200)

    with st.spinner("Checking SSL Certificate..."):
        st.session_state.ssl_result = check_ssl(target)
    st.text_area("SSL Info", st.session_state.ssl_result, height=100)

    with st.spinner("Checking HTTP Security Headers..."):
        st.session_state.header_result = check_http_headers(target)
    st.text_area("HTTP Headers", st.session_state.header_result, height=150)

    with st.spinner("Analyzing Findings..."):
        st.session_state.analysis = simple_vulnerability_analysis(
            st.session_state.nmap_result,
            st.session_state.nikto_result,
            st.session_state.header_result
        )
    st.text_area("Vulnerability Analysis", st.session_state.analysis, height=200)

if (
    "nmap_result" in st.session_state and
    "nikto_result" in st.session_state and
    "ssl_result" in st.session_state and
    "header_result" in st.session_state and
    "analysis" in st.session_state
):
    if st.button("Generate DOCX Report"):
        docx_file = save_report(
            st.session_state.nmap_result,
            st.session_state.nikto_result,
            st.session_state.ssl_result,
            st.session_state.header_result,
            st.session_state.analysis
        )
        if docx_file:
            st.download_button(
                label="Download Audit Report",
                data=docx_file,
                file_name="audit_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            st.error("The report could not be generated.")
